#!/usr/bin/env python3
from __future__ import annotations

import html
import json
import re
import time
from dataclasses import dataclass, asdict
from pathlib import Path
from urllib.parse import quote
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html as lxml_html


OUTPUT_DIR = Path("output/alina_job_manual")
CITIES = ["重庆", "成都", "苏州", "上海", "杭州"]
PLATFORM_POLICY = {
    "智联招聘": {
        "status": "可自动抓取",
        "reason": "公开搜索页 HTML 内含岗位卡片，可解析标题、公司、薪资、地点、经验、学历、标签、链接。",
        "action": "作为默认数据源；抓取后必须缓存原始 HTML，并抽查高分和低分岗位。",
    },
    "BOSS直聘": {
        "status": "需要登录态浏览器",
        "reason": "页面主要由前端动态加载，公开 HTML 只有应用壳，直接抓取会拿不到岗位卡片。",
        "action": "不混入自动清单；如用户明确开启浏览器调试并登录，再用浏览器可见页面抽取。",
    },
    "拉勾招聘": {
        "status": "需要登录态浏览器",
        "reason": "搜索结果依赖动态请求和风控，公开页面不稳定。",
        "action": "只作为备用入口；未验证岗位卡片前不入库。",
    },
    "前程无忧": {
        "status": "需要动态渲染验证",
        "reason": "公开 HTML 多为前端应用壳，岗位数据不稳定内嵌。",
        "action": "只作为备用入口；需要浏览器渲染后确认能读到岗位卡片。",
    },
    "猎聘": {
        "status": "需要动态渲染验证",
        "reason": "公开搜索页容易返回 SEO/壳页面，和用户看到的岗位列表不一致。",
        "action": "只作为备用入口；未验证可读岗位卡片前不入库。",
    },
    "应届生求职网": {
        "status": "旧入口失效",
        "reason": "本次链接库里的旧搜索链接返回 404。",
        "action": "不自动抓取旧链接；后续手动改用站内新搜索入口。",
    },
}
KEYWORDS = [
    "AI产品运营",
    "AI应用运营",
    "AIGC内容运营",
    "大模型产品助理",
    "数据运营",
    "用户增长运营",
    "内容运营",
    "新媒体运营",
    "产品经理助理",
    "前端开发实习",
    "小程序开发",
    "Node.js开发",
]

STRONG_TERMS = [
    "AI",
    "AIGC",
    "大模型",
    "Prompt",
    "产品运营",
    "内容运营",
    "数据运营",
    "用户增长",
    "产品助理",
    "产品经理助理",
    "数据分析",
    "用户运营",
    "新媒体",
    "小程序",
    "前端",
    "Node",
    "人工智能",
]

ENTRY_TERMS = ["经验不限", "应届", "实习", "校招", "1年以下", "专人带教"]
EXCLUDE_TERMS = [
    "电话销售",
    "电销",
    "无责",
    "招商",
    "加盟",
    "地推",
    "房产",
    "保险",
    "贷款",
    "客服",
    "销售顾问",
    "网络销售",
    "课程顾问",
    "主播",
    "带货",
    "骑手",
    "送餐",
    "司机",
    "驾驶员",
    "普工",
    "操作工",
    "咨询师",
    "外贸业务",
    "销售助理",
    "助理销售",
]
SENIOR_TERMS = ["总监", "负责人", "高级", "专家", "主管", "经理（", "经理岗", "5-10年", "10年以上"]
DEV_EXCLUDE_TERMS = ["Java", "C++", "嵌入式", "硬件", "测试开发", "运维", "算法工程师"]
STRICT_SENIOR_TITLE_TERMS = ["总监", "负责人", "资深", "专家", "主管", "组长"]
TITLE_DIRECTION_TERMS = [
    "AI",
    "ai",
    "AIGC",
    "大模型",
    "模型",
    "数据运营",
    "运营数据",
    "数据服务运营",
    "内容运营",
    "新媒体运营",
    "产品运营",
    "用户运营",
    "用户增长",
    "产品经理助理",
    "产品助理",
    "数据标注",
    "前端开发实习",
    "前端实习",
    "小程序",
    "Node.js",
    "node.js",
]


@dataclass
class Job:
    city: str
    keyword: str
    title: str
    company: str
    salary: str
    location: str
    experience: str
    degree: str
    tags: list[str]
    url: str
    score: int
    reasons: list[str]
    source: str = "智联招聘"


def fetch_url(url: str) -> str:
    req = Request(
        url,
        headers={
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0 Safari/537.36"
            ),
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        },
    )
    with urlopen(req, timeout=25) as resp:
        data = resp.read()
    return data.decode("utf-8", errors="ignore")


def set_east_asia_font(run, font_name: str = "微软雅黑") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def clean_text(parts: list[str] | str) -> str:
    if isinstance(parts, str):
        value = parts
    else:
        value = " ".join(p.strip() for p in parts if str(p).strip())
    return re.sub(r"\s+", " ", value).strip()


def parse_zhaopin(html_text: str, city: str, keyword: str) -> list[dict]:
    tree = lxml_html.fromstring(html_text)
    items = tree.xpath(
        '//*[contains(concat(" ", normalize-space(@class), " "), " joblist-box__item ")]'
    )
    jobs = []
    for item in items:
        title = clean_text(item.xpath('.//a[contains(@class,"jobinfo__name")]/text()'))
        url = clean_text(item.xpath('.//a[contains(@class,"jobinfo__name")]/@href'))
        salary = clean_text(item.xpath('.//*[contains(@class,"jobinfo__salary")]/text()'))
        company = clean_text(item.xpath('.//*[contains(@class,"companyinfo__name")]/text()'))
        other = [
            clean_text(x)
            for x in item.xpath('.//*[contains(@class,"jobinfo__other-info")]/descendant::text()')
            if clean_text(x)
        ]
        tags = [
            clean_text(x)
            for x in item.xpath('.//*[contains(@class,"joblist-box__item-tag")]/text()')
            if clean_text(x)
        ]
        if not title or not url:
            continue
        jobs.append(
            {
                "city": city,
                "keyword": keyword,
                "title": title,
                "company": company,
                "salary": salary,
                "location": other[0] if len(other) > 0 else city,
                "experience": other[1] if len(other) > 1 else "",
                "degree": other[2] if len(other) > 2 else "",
                "tags": tags,
                "url": url,
            }
        )
    return jobs


def score_job(raw: dict) -> tuple[int, list[str], bool]:
    text = " ".join(
        [
            raw.get("title", ""),
            raw.get("company", ""),
            raw.get("salary", ""),
            raw.get("location", ""),
            raw.get("experience", ""),
            raw.get("degree", ""),
            " ".join(raw.get("tags", [])),
        ]
    )
    lower = text.lower()
    title = raw.get("title", "")
    reasons: list[str] = []

    if any(term.lower() in lower for term in EXCLUDE_TERMS):
        return 0, ["排除：偏销售/客服/地推/招商"], False
    if "销售" in title and "销售运营" not in title:
        return 0, ["排除：偏销售"], False
    if any(term in title for term in STRICT_SENIOR_TITLE_TERMS):
        return 0, ["排除：偏资深"], False
    if any(term in raw.get("experience", "") for term in ["5-10年", "10年以上"]):
        return 0, ["排除：经验要求过高"], False
    if any(term.lower() in title.lower() for term in DEV_EXCLUDE_TERMS) and not any(
        term in title for term in ["前端", "小程序", "Node"]
    ):
        return 0, ["排除：开发方向不匹配"], False
    title_matches_direction = any(term in title for term in TITLE_DIRECTION_TERMS)
    if not title_matches_direction:
        return 0, ["排除：标题方向不匹配"], False
    if any(term in title for term in ["前端", "小程序", "Node", "node"]) and "实习" not in title and raw.get("experience") not in ["经验不限", "1年以下"]:
        return 0, ["排除：开发经验要求偏高"], False

    score = 10
    if raw["city"] in raw.get("location", "") or raw["city"] in text:
        score += 10
        reasons.append(f"城市匹配：{raw['city']}")

    matched = []
    for term in STRONG_TERMS:
        if term.lower() in lower:
            matched.append(term)
            score += 6
    for term in matched[:4]:
        reasons.append(f"命中：{term}")

    title_lower = title.lower()
    keyword = raw.get("keyword", "")
    if keyword and keyword.lower() in title_lower:
        score += 12
        reasons.append("标题直匹配")
    elif any(part.lower() in title_lower for part in ["ai", "aigc", "大模型", "数据运营", "内容运营", "新媒体", "产品运营", "前端", "小程序", "node"]):
        score += 8
        reasons.append("标题方向匹配")

    if any(term in text for term in ENTRY_TERMS):
        score += 10
        reasons.append("新人友好")
    if any(term in text for term in SENIOR_TERMS):
        score -= 15
        reasons.append("注意：可能偏资深")

    if "150-" in raw.get("salary", "") or "200" in raw.get("salary", ""):
        score += 4
    if re.search(r"[6-9]000-[8-9]000|[7-9]000|1[0-9]000", raw.get("salary", "")):
        score += 6
        reasons.append("薪资可关注")

    score = max(0, min(100, score))
    return score, reasons[:5], bool(matched) and score >= 42


def crawl_zhaopin() -> list[Job]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    raw_dir = OUTPUT_DIR / "raw_pages"
    raw_dir.mkdir(exist_ok=True)
    jobs: dict[str, Job] = {}
    failures = []

    for city in CITIES:
        for keyword in KEYWORDS:
            query = f"{city} {keyword}"
            url = f"https://sou.zhaopin.com/?kw={quote(query)}"
            cache_path = raw_dir / f"{city}_{keyword}.html"
            fetched = False
            try:
                if cache_path.exists():
                    html_text = cache_path.read_text(encoding="utf-8", errors="ignore")
                else:
                    html_text = fetch_url(url)
                    cache_path.write_text(html_text, encoding="utf-8")
                    fetched = True
                parsed = parse_zhaopin(html_text, city, keyword)
            except Exception as exc:
                failures.append({"city": city, "keyword": keyword, "error": str(exc)})
                continue

            for raw in parsed:
                score, reasons, keep = score_job(raw)
                if not keep:
                    continue
                key = raw["url"]
                job = Job(score=score, reasons=reasons, **raw)
                old = jobs.get(key)
                if old is None or job.score > old.score:
                    jobs[key] = job
            if fetched:
                time.sleep(0.6)

    result = sorted(jobs.values(), key=lambda j: (j.score, j.city), reverse=True)
    (OUTPUT_DIR / "zhaopin_jobs_scored.json").write_text(
        json.dumps([asdict(j) for j in result], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (OUTPUT_DIR / "crawl_failures.json").write_text(
        json.dumps(failures, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_platform_policy()
    return result


def write_platform_policy() -> Path:
    path = OUTPUT_DIR / "platform_access_policy.json"
    path.write_text(json.dumps(PLATFORM_POLICY, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def grouped_jobs(jobs: list[Job]) -> dict[str, list[Job]]:
    groups = {city: [] for city in CITIES}
    for job in jobs:
        groups.setdefault(job.city, []).append(job)
    for city in groups:
        groups[city].sort(key=lambda j: j.score, reverse=True)
    return groups


def build_html(jobs: list[Job]) -> Path:
    groups = grouped_jobs(jobs)
    out = OUTPUT_DIR / "张梦玲-高薪岗位投递手册.html"
    rows = []
    for job in jobs:
        rows.append(
            f"""
<tr>
  <td>{job.score}</td>
  <td>{html.escape(job.city)}</td>
  <td><a href="{html.escape(job.url)}">{html.escape(job.title)}</a><br><span>{html.escape(job.company)}</span></td>
  <td>{html.escape(job.salary)}</td>
  <td>{html.escape(job.location)}<br><span>{html.escape(job.experience)} / {html.escape(job.degree)}</span></td>
  <td>{html.escape("；".join(job.reasons))}</td>
</tr>"""
        )
    city_cards = "\n".join(
        f"<div class='card'><b>{city}</b><strong>{len(groups.get(city, []))}</strong><span>个可投岗位</span></div>"
        for city in CITIES
    )
    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>张梦玲高薪岗位投递手册</title>
<style>
body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC",Arial,sans-serif;margin:0;background:#f6f7f9;color:#172033;}}
main{{max-width:1180px;margin:0 auto;padding:36px 28px 60px;}}
h1{{font-size:34px;margin:0 0 8px;}}
h2{{margin-top:34px;border-bottom:2px solid #d8e0ea;padding-bottom:8px;}}
p{{line-height:1.7;color:#4a5568;}}
.hero{{background:white;border:1px solid #dfe5ed;border-radius:8px;padding:28px;}}
.grid{{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin-top:22px;}}
.card{{background:#eef4f8;border:1px solid #d8e2ec;border-radius:8px;padding:14px;}}
.card strong{{display:block;font-size:30px;margin-top:8px;color:#175d7a;}}
.note{{background:#fff7e6;border:1px solid #efd7a0;border-radius:8px;padding:14px 18px;margin-top:18px;}}
table{{width:100%;border-collapse:collapse;background:white;border:1px solid #dfe5ed;margin-top:16px;}}
th,td{{border-bottom:1px solid #e7ebf0;padding:10px 12px;text-align:left;vertical-align:top;font-size:14px;}}
th{{background:#edf3f7;color:#17384f;}}
td:first-child{{font-weight:700;color:#175d7a;}}
a{{color:#0b65c2;text-decoration:none;}}
span{{color:#6b7280;font-size:12px;}}
ol li{{margin:7px 0;}}
</style>
</head>
<body>
<main>
<section class="hero">
  <h1>张梦玲 · 高薪岗位投递手册</h1>
  <p>依据你的简历方向筛选：AI产品运营、AI应用运营、AIGC内容运营、数据运营、用户增长、产品助理、前端/小程序/Node 入门岗位。数据来源：投递链接库里的智联招聘公开搜索结果，生成日期：2026-06-16。</p>
  <div class="grid">{city_cards}</div>
  <div class="note">先投 80 分以上，再投 65 分以上；偏销售、客服、招商、地推、明显资深和重开发岗位已剔除。</div>
</section>
<h2>投递顺序</h2>
<ol>
  <li>先投 AI产品运营 / AIGC内容运营 / AI应用运营，这三类最贴合你的 AI 工具、内容和数据经历。</li>
  <li>再投数据运营、用户增长运营、产品经理助理，突出 Excel/SQL、指标拆解、A/B 测试、PRD 和项目推进。</li>
  <li>前端/小程序/Node 只投实习、初级、项目助理型岗位，避免纯重开发岗位。</li>
</ol>
<h2>可投岗位清单（{len(jobs)} 个）</h2>
<table>
<thead><tr><th>分数</th><th>城市</th><th>岗位/公司</th><th>薪资</th><th>地点/要求</th><th>为什么适合</th></tr></thead>
<tbody>{''.join(rows)}</tbody>
</table>
<h2>反爬与受限平台处理</h2>
<p>本手册只收录已验证能从公开页面读取岗位卡片的智联招聘结果。BOSS、拉勾、前程无忧、猎聘等平台需要登录态浏览器或动态渲染验证，未验证前只作为备用入口，不把不可验证岗位混进清单。</p>
<table>
<thead><tr><th>平台</th><th>状态</th><th>处理方式</th></tr></thead>
<tbody>{''.join(f"<tr><td>{html.escape(name)}</td><td>{html.escape(info['status'])}</td><td>{html.escape(info['action'])}</td></tr>" for name, info in PLATFORM_POLICY.items())}</tbody>
</table>
</main>
</body>
</html>"""
    out.write_text(html_text, encoding="utf-8")
    return out


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    set_east_asia_font(run)
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)
    tbl_pr.append(borders)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B65C2")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 2" else 18)
        style.paragraph_format.space_after = Pt(7 if style_name == "Heading 2" else 10)


def build_docx(jobs: list[Job]) -> Path:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("张梦玲 · 高薪岗位投递手册")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    set_east_asia_font(run)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("生成日期：2026-06-16｜来源：投递链接库 + 智联招聘公开岗位")

    doc.add_heading("一、先投什么", level=1)
    for text in [
        "第一梯队：AI产品运营、AI应用运营、AIGC内容运营、大模型产品助理。",
        "第二梯队：数据运营、用户增长运营、内容运营、产品经理助理。",
        "补充投递：前端开发实习、小程序开发、Node.js开发，只选实习或初级岗位。",
        "已剔除：明显销售、客服、招商、地推、重资深、重开发岗位。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    groups = grouped_jobs(jobs)
    doc.add_heading("二、城市概览", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    headers = ["城市", "可投岗位数", "建议动作"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True)
        shade_cell(cell, "E8EEF5")
    for city in CITIES:
        row = table.add_row().cells
        set_cell_text(row[0], city)
        set_cell_text(row[1], str(len(groups.get(city, []))))
        set_cell_text(row[2], "先投80分以上，再补65分以上")

    doc.add_heading(f"三、可投岗位清单（{len(jobs)}个）", level=1)
    table = doc.add_table(rows=1, cols=7)
    set_table_borders(table)
    widths = [0.45, 0.55, 1.45, 1.45, 0.8, 1.65, 0.75]
    headers = ["分", "城", "岗位", "公司", "薪资", "匹配理由", "链接"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, headers[idx], bold=True)
        shade_cell(cell, "E8EEF5")
        cell.width = Inches(widths[idx])

    for job in jobs:
        row = table.add_row().cells
        values = [
            str(job.score),
            job.city,
            job.title,
            job.company,
            job.salary,
            "；".join(job.reasons),
            "",
        ]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)
            row[idx].width = Inches(widths[idx])
        add_hyperlink(row[6].paragraphs[0], "打开", job.url)

    doc.add_heading("四、备用入口说明", level=1)
    doc.add_paragraph(
        "BOSS、拉勾、前程无忧、猎聘页面主要依赖动态加载或风控；当前没有使用你的浏览器登录态，所以不把无法验证的岗位混进清单。"
    )
    doc.add_paragraph("应届生求职网原搜索链接返回404，建议手动改用站内新搜索。")
    doc.add_paragraph("反爬处理原则：公开页面能读到岗位卡片才入库；需要登录态或动态渲染的平台只做备用入口。")

    out = OUTPUT_DIR / "张梦玲-高薪岗位投递手册.docx"
    doc.save(out)
    return out


def main() -> None:
    jobs = crawl_zhaopin()
    html_path = build_html(jobs)
    docx_path = build_docx(jobs)
    summary = {
        "jobs": len(jobs),
        "html": str(html_path.resolve()),
        "docx": str(docx_path.resolve()),
        "platform_policy": str((OUTPUT_DIR / "platform_access_policy.json").resolve()),
        "by_city": {city: len(grouped_jobs(jobs).get(city, [])) for city in CITIES},
    }
    (OUTPUT_DIR / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
